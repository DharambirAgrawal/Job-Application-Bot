import io
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
import PyPDF2

class DocumentUtils:
    """Utility class for working with DOCX and PDF files."""

    @staticmethod
    def update_docx_placeholders(doc_source, replacements):
        """
        Replace placeholders in a DOCX document without changing formatting.

        Args:
            doc_source: path to DOCX file, bytes, or file-like object
            replacements: dict of placeholders -> replacement text

        Returns:
            BytesIO with updated document
        """
        if isinstance(doc_source, (str, Path)):
            doc_stream = open(doc_source, "rb")
        elif hasattr(doc_source, "read"):
            doc_stream = doc_source
        else:
            doc_stream = io.BytesIO(doc_source)

        doc = Document(doc_stream)
        DocumentUtils._replace_in_paragraphs(doc.paragraphs, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    DocumentUtils._replace_in_paragraphs(cell.paragraphs, replacements)

        for section in doc.sections:
            DocumentUtils._replace_in_paragraphs(section.header.paragraphs, replacements)
            DocumentUtils._replace_in_paragraphs(section.footer.paragraphs, replacements)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    @staticmethod
    def _replace_in_paragraphs(paragraphs, replacements):
        for paragraph in paragraphs:
            for old, new in replacements.items():
                if not old:
                    continue
                DocumentUtils._replace_text(paragraph, old, new)

    # @staticmethod
    # def _replace_text(paragraph, old, new):
    #     """Replace text inside paragraph runs without changing formatting."""
    #     replaced = False
    #     for run in paragraph.runs:
    #         if old in run.text:
    #             run.text = run.text.replace(old, new)
    #             replaced = True
    #     if replaced:
    #         return

    #     # Handle placeholders spanning multiple runs
    #     while True:
    #         text_combined = "".join(run.text for run in paragraph.runs)
    #         start = text_combined.find(old)
    #         if start == -1:
    #             break
    #         end = start + len(old)
    #         DocumentUtils._replace_across_runs(paragraph, start, end, new)
    
    @staticmethod
    def _replace_text(paragraph, old, new):
        """Replace text inside paragraph runs without changing formatting.
        If new is empty, remove the placeholder. If the paragraph becomes empty, remove it completely.
        """
        if old not in paragraph.text:
            return

        # If replacement is empty, remove the placeholder
        if new == "":
            for run in paragraph.runs:
                run.text = run.text.replace(old, "")
            # If paragraph becomes empty, clear all runs
            if paragraph.text.strip() == "":
                for run in paragraph.runs:
                    run.text = ""
            return

        # Normal replacement
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)

        # Handle placeholders spanning multiple runs
        while True:
            combined = "".join(run.text for run in paragraph.runs)
            start = combined.find(old)
            if start == -1:
                break
            end = start + len(old)
            DocumentUtils._replace_across_runs(paragraph, start, end, new)


    @staticmethod
    def _replace_across_runs(paragraph, start, end, replacement):
        run_positions = []
        pos = 0
        for run in paragraph.runs:
            run_positions.append((run, pos, pos + len(run.text)))
            pos += len(run.text)

        start_idx = end_idx = None
        for i, (run, s, e) in enumerate(run_positions):
            if s <= start < e:
                start_idx = i
            if s <= end - 1 < e:
                end_idx = i
        if start_idx is None or end_idx is None:
            return

        start_run, s_start, s_end = run_positions[start_idx]
        end_run, e_start, e_end = run_positions[end_idx]

        start_offset = start - s_start
        end_offset = end - e_start

        prefix = start_run.text[:start_offset]
        suffix = end_run.text[end_offset:]

        start_run.text = prefix + replacement
        for i in range(start_idx + 1, end_idx):
            paragraph.runs[i].text = ""
        end_run.text = suffix

    @staticmethod
    def extract_text(file_source):
        """
        Extract plain text from DOCX or PDF.

        Args:
            file_source: path to file (str/Path) or file-like object

        Returns:
            string with extracted text
        """
        # If file_source is a path
        if isinstance(file_source, (str, Path)):
            path = Path(file_source)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_source}")

            suffix = path.suffix.lower()
            if suffix == ".docx":
                doc = Document(path)
                text_parts = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_parts.append(cell.text)
                return "\n".join(text_parts)
            elif suffix == ".pdf":
                text = []
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text.append(page.extract_text() or "")
                return "\n".join(text)
            else:
                raise ValueError("Unsupported file type. Only DOCX and PDF allowed.")

        # If file_source is bytes or file-like object
        elif hasattr(file_source, "read"):
            start_pos = file_source.tell()
            content = file_source.read()
            file_source.seek(start_pos)

            # Try DOCX first
            try:
                doc = Document(io.BytesIO(content))
                text_parts = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_parts.append(cell.text)
                return "\n".join(text_parts)
            except Exception:
                # Try PDF
                try:
                    pdf_stream = io.BytesIO(content)
                    reader = PyPDF2.PdfReader(pdf_stream)
                    text = []
                    for page in reader.pages:
                        text.append(page.extract_text() or "")
                    return "\n".join(text)
                except Exception:
                    raise ValueError("Unsupported file format or failed to read the file.")

        else:
            raise TypeError("file_source must be a path, bytes, or file-like object")

"""
# Update a DOCX
replacements = {
    "<%DATE%>": "2025-10-20",
    "<%TEAM%>": "Engineering",
    "<%COMPANYNAME%>": "OpenAI",
    "<%LOCATION%>": "San Francisco",
    "<%GREETINGS%>": "Dear Hiring Manager",
    "<%PARAGRAPH%>": "I am excited to apply for this role..."
}

updated_doc = DocumentUtils.update_docx_placeholders("template.docx", replacements)
with open("output.docx", "wb") as f:
    f.write(updated_doc.read())

# Extract text from DOCX or PDF
text = DocumentUtils.extract_text("resume.pdf")
print(text)


"""